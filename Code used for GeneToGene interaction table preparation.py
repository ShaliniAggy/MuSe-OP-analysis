#!/usr/bin/env python
# coding: utf-8

# In[656]:


from docx import Document
import os
# Specify the path to your Word file
file_path = r'C:\Users\CHANIR.WISMAIN\Downloads\activates, 3000genes.docx'

##r'C:\Users\CHANIR.WISMAIN\Downloads\Activates, B file best 100 product BC genes word.docx'

#Binds to, B file best 100 product BC genes word.docx
#Interacts, B file best 100 product BC genes word.docx
#C:\Users\CHANIR.WISMAIN\Downloads\Activates, B file best 100 product BC genes word.docx
# Open the Word document
doc = Document(file_path)
# extract the search query 
filename = os.path.basename(file_path)
name_until_first_comma = filename.split(',')[0]
name_until_first_comma_lower = name_until_first_comma.lower().strip()
#print(name_until_first_comma)
#Access paragraphs in the document
data =[]
for paragraph in doc.paragraphs:
    if not paragraph.text.startswith('QUERY: '):
        data.append(paragraph.text)
        #print(paragraph.text)


# In[657]:


import re

def remove_html_tags_from_list(html_list):
    clean_list = [re.sub('<.*?>', '', item) for item in html_list]
    return clean_list



# Remove HTML tags from the list
cleaned_tags_list = remove_html_tags_from_list(data)

# Print the cleaned list
for item in cleaned_tags_list:
    print(item)


# In[659]:


import pandas as pd

#extract the gene list from gene file
Gene_list_file = r'C:\Users\CHANIR.WISMAIN\Downloads\3000genes.csv'
#r'C:\Users\CHANIR.WISMAIN\Downloads\Copy of top 1000_functional score_lncRNA_Disease_MuSe_Input.csv'
#r'C:\Users\CHANIR.WISMAIN\Downloads\List_lncRNA_GIFtS_Disorders_3000_SA_08012023.csv'
#r'C:\Users\CHANIR.WISMAIN\Downloads\B file best 100 product BC genes text.txt'
#Gene_list = pd.read_csv(Gene_list_file, header=None).squeeze("columns")
Gene_list  = pd.read_csv(Gene_list_file)
Gene_list = Gene_list['gene'].tolist()



# In[660]:


from docx import Document
import re
# Specify the path to your Word file
#file_path = r'C:\Users\CHANIR.WISMAIN\Downloads\Activates, B file best 100 product BC genes word.docx'

# Strings you want to count occurrences for
genes_list = Gene_list 
# Dictionary to store gene counts
gene_counts = {gene: 0 for gene in genes_list}

# Open the Word document
doc = cleaned_tags_list
#Document(file_path)

# Flag to ignore lines with 'QUERY:'
ignore_query = False

# Iterate through paragraphs
for paragraph in doc:
    text = paragraph.strip()
    # Replace '(' and ')' with an empty string
    text = text.replace('(', '').replace(')', '')

    if text.startswith('QUERY:'):
        # Set flag to ignore lines with 'QUERY:'
        ignore_query = True
    elif text:
        if ignore_query:
            # Reset flag after encountering 'QUERY:'
            ignore_query = False
        else:
            # Check for gene presence and count
            for gene in genes_list:
                
                #name_until_first_comma_lower
                if name_until_first_comma_lower in text and gene in text :
                    gene_counts[gene] += 1
# Display gene counts
for gene, count in gene_counts.items():
    print(f"{gene}: {count} times")


# In[661]:


# Create a new list of genes with counts > 0
selected_genes = [gene for gene, count in gene_counts.items() if count > 0]


# In[662]:


from docx import Document
# Set to store unique sentences
unique_sentences = set()

gene_dict = {}
# Open the Word document
doc = cleaned_tags_list

# List of genes
genes_list = selected_genes

# Flag to ignore lines with 'QUERY:'
ignore_query = False

# Iterate through paragraphs
for paragraph in doc:
    text = paragraph.strip()
    # Replace '(' and ')' with an empty string
    text = text.replace('(', '').replace(')', '')
    text = text.replace(',', '').replace('...', '').replace('.', '')
    if text.startswith('QUERY:'):
        # Set flag to ignore lines with 'QUERY:'
        ignore_query = True
    elif text:
        if ignore_query:
            # Reset flag after encountering 'QUERY:'
            ignore_query = False
        else:
            # Check for 'activates' and add unique sentences to the set
            if name_until_first_comma_lower in text:
                text1 = text[:text.find(name_until_first_comma_lower)]
                # Strip everything before the gene name
                for gene in genes_list:
                    if gene in text:
                        #stripped_sentence = text[text.find(name_until_first_comma_lower)-(len(gene)+1):]
                        stripped_sentence_temp = text[text.find(name_until_first_comma_lower)+len(name_until_first_comma_lower):]
                        # Split the stripped_sentence_temp into words
                        words = stripped_sentence_temp.split()

                        # Take only the first 7 words
                        stripped_sentence_temp = ' '.join(words[:5])
                        # Take only the first 7 words
                        stripped_sentence_temp = ' '.join(words[:5])
                        unique_sentences.add(stripped_sentence_temp)
                        #gene_dict[gene].append(stripped_sentence_temp)
                        #for the join method
                        gene_dict.setdefault(gene, []).append(stripped_sentence_temp)

# count the number od wordes for striping -len of search query + gene+ 10.
len_sequence = len(name_until_first_comma.split())

for sentence in unique_sentences:
    words = sentence.split()
    if words:
        words[0] = words[0].upper()
        truncated_sentence = ' '.join(words[:5])
        #print(truncated_sentence)


# Convert the dictionary to a pandas DataFrame
gene_sentence_dataframe = pd.DataFrame(list(gene_dict.items()), columns=['Gene', 'sentence'])
gene_sentence_dataframe['sentence'] = gene_sentence_dataframe['sentence'].apply(lambda x: set(x))
gene_sentence_dataframe = gene_sentence_dataframe.explode('sentence')


display(gene_sentence_dataframe )
print(len(unique_sentences))


# In[663]:


file_path = r'C:\Users\CHANIR.WISMAIN\Downloads\Alias_genes.csv'
file_path1 = r'C:\Users\CHANIR.WISMAIN\Downloads\dorons_pj.csv'
df_db_gene_info = pd.read_csv(file_path)
df_db_gene_info_gene = pd.read_csv(file_path1)
#df_db_gene_info.columns = ['symbol','name','Alias Name','Type','category','class']
#df_db_gene_info= df_db_gene_info[['symbol','name','Alias Name','category','class']]
df_db_gene_info.columns = ['symbol','name','category','class']
df_db_gene_info_gene.columns = ['symbol','name','category','class']
selected_columns_alias = ['symbol', 'name', 'category', 'class']
selected_columns = ['symbol', 'name', 'category', 'class']
df_db_gene_info1 = df_db_gene_info_gene[selected_columns]
df_db_gene_info_alias= df_db_gene_info[selected_columns_alias]
df_db_gene_info_alias.columns=['symbol', 'name', 'category', 'class']
df_db_gene_info1


# In[664]:


gene_list_l1 = [gene for gene in df_db_gene_info1['symbol'].tolist()]
gene_list_l1 = [str(gene).upper() for gene in gene_list_l1]
df_db_gene_info = df_db_gene_info[pd.isnull(df_db_gene_info['symbol'])]

gene_list_l2 = [gene for gene in df_db_gene_info_alias['symbol'].tolist()]
gene_list_l2 = [str(gene).upper().replace('-', '') for gene in gene_list_l2]
gene_list_l = gene_list_l1 + gene_list_l2

gene_list_l = list(set(gene_list_l))


# In[665]:


# Lists to store genes, sentences, and matched words
# Dictionary to store genes, sentences, and matched words
gene_word_dict = {'Gene': [], 'sentence': [], 'words_list': []}
genes = []
sentences = []
words_list = []
data_list = []
# Iterate through each row in the DataFrame
for index, row in gene_sentence_dataframe.iterrows():
    data_list = []
    gene = row['Gene']
    stripped_sentence = row['sentence']
    genes.append(gene)
    sentences.append(stripped_sentence)
    for word in stripped_sentence.split():
        if word.upper() in gene_list_l:
            data_list.append(word.upper())
        elif word.upper().replace('-', '') in gene_list_l:
            data_list.append(word.upper().replace('-', ''))
        elif word.upper().replace('-', ' ').split()[0] in gene_list_l:
            data_list.append(word.upper().replace('-', ' ').split()[0])
      
    words_list.append([gene,stripped_sentence,list(set(data_list))])
         
    


# In[667]:


# Create a DataFrame from the dictionary
gene_word_dataframe = pd.DataFrame(words_list)
gene_word_dataframe.columns = ['Gene', 'sentence','words_list']
# Combine rows with the same 'Gene' and 'sentence' values and create a list of 'Gene' values
# Group by 'Gene' and 'sentence', and aggregate the 'words_list' as a list
# Group by 'Gene' and 'sentence', and aggregate the 'words_list' as a list
combined_df = gene_word_dataframe.groupby(['Gene', 'sentence']).agg({'words_list': 'sum'}).reset_index()
combined_df.columns = ['Gene', 'sentence', 'combined_genes']
# Filter out lines containing 'lncRNA'
combined_df = combined_df[~combined_df['Gene'].str.contains('lncRNA')]
# Filter out rows with empty 'combined_genes' list
filtered_df = combined_df[combined_df['combined_genes'].apply(len) > 0]

# Display the resulting DataFrame
print(filtered_df)
print(len(filtered_df))
print(len(combined_df))


# In[668]:


# Reset the index to make it unique
filtered_df = filtered_df.reset_index(drop=True)

# Use explode() on 'combined_genes' column
filtered_df = filtered_df.explode('combined_genes')

filtered_df = filtered_df[~filtered_df['combined_genes'].str.contains('TO')]
filtered_df = filtered_df[~filtered_df['combined_genes'].str.contains('OF')]
filtered_df = filtered_df[~filtered_df['combined_genes'].str.contains('CELL')]
filtered_df = filtered_df[~filtered_df['combined_genes'].str.contains('IN')]
filtered_df = filtered_df[~filtered_df['combined_genes'].str.contains('AS')]
filtered_df = filtered_df[~filtered_df['combined_genes'].str.contains('FOR')]
filtered_df = filtered_df[~filtered_df['combined_genes'].str.contains('AN')]
#filtered_df = filtered_df[~filtered_df['combined_genes'].str.contains('MIR')]
filtered_df = filtered_df[filtered_df['combined_genes'].str.len() >= 2]
filtered_df = filtered_df[~filtered_df['combined_genes'].str.isnumeric()]
filtered_df = filtered_df[filtered_df['combined_genes'] != 'MRI']
filtered_df = filtered_df[filtered_df['combined_genes'] != 'VIA']
filtered_df = filtered_df[filtered_df['combined_genes'] != 'UP']
filtered_df = filtered_df[filtered_df['combined_genes'] != 'AT']
filtered_df = filtered_df[~filtered_df['combined_genes'].isin(['MIR', 'VIA', 'UP', 'AT'])]


# In[670]:


# Assuming 'exploded_df' is your DataFrame with columns 'Gene', 'sentence', and 'combined_genes'

# Truncate the 'sentence' column to contain only the first 7 words for each row
filtered_df['sentence'] = filtered_df['sentence'].apply(lambda x: ' '.join(x.split()[:5]))

# Display the resulting DataFrame
print(filtered_df)


# In[673]:


# Assuming df_db_gene_info1 and df_db_gene_info_alias are your DataFrames
common_symbols = set(df_db_gene_info1['symbol'])
filtered_df_db_gene_info_alias = df_db_gene_info_alias[~df_db_gene_info_alias['symbol'].isin(common_symbols)]

df_db_gene_info11 = pd.concat([df_db_gene_info1, filtered_df_db_gene_info_alias])
df_db_gene_info11


# In[676]:


# Create a new dataframe by merging df1 and df2 on the 'Gene' and 'symbol' columns
merged_df_word_db = pd.merge(filtered_df, df_db_gene_info11, left_on='combined_genes', right_on='symbol', how='inner')
merged_df_word_db=merged_df_word_db.drop_duplicates()


# In[677]:


sub_df_db_gene_info1=df_db_gene_info11[['symbol','name']]

# Assuming df is your DataFrame
sub_df_db_gene_info1 = sub_df_db_gene_info1.dropna(subset=['name'])

# If you want to reset the index after dropping rows
sub_df_db_gene_info1 = sub_df_db_gene_info1.reset_index(drop=True)
len(sub_df_db_gene_info1)


# In[678]:


sub_df_db_gene_info1=df_db_gene_info11[['symbol','name']]
# Create a new dataframe by merging df1 and df2 on the 'Gene' and 'symbol' columns
merged_df_word_db_final = pd.merge(merged_df_word_db, sub_df_db_gene_info1 ,left_on='Gene', right_on='symbol', how='inner')
merged_df_word_db_final


# In[679]:


merged_df_word_db_final = merged_df_word_db_final[['Gene','name_y','combined_genes','sentence','category','class','name_x']]
#merged_df_word_db_final.columns['imput Gene symbol','Gene description','Target Gene symbol','7 words for the keywords','catagory','class','Description of the target Gene']
merged_df_word_db_final.columns = ['Imput Gene symbol', 'Gene description', 'Target Gene symbol', '5 words for the keywords', 'catagory', 'class', 'Description of the target Gene']


# In[680]:


keyword = [name_until_first_comma_lower for x in merged_df_word_db_final['Imput Gene symbol']]
merged_df_word_db_final['Keyword']=keyword
merged_df_word_db_final=merged_df_word_db_final[['Imput Gene symbol', 'Gene description','Keyword', 'Target Gene symbol', '5 words for the keywords', 'catagory', 'class', 'Description of the target Gene']]


# In[681]:


# Filter rows where 'Imput Gene symbol' is not equal to 'Target Gene symbol'
filtered_df = merged_df_word_db_final[merged_df_word_db_final['Imput Gene symbol'] != merged_df_word_db_final['Target Gene symbol']]
filtered_df = filtered_df[filtered_df['Imput Gene symbol'] != filtered_df['Gene description']]

filtered_df = filtered_df.drop_duplicates()


# In[683]:


#gene_word_dataframe.to_csv('gene_word_dataframe.txt')
# Specify the dictionary name in the file path
file_path = r'C:\Users\CHANIR.WISMAIN\Downloads\gene_word_activates4.csv'
#r'C:\Users\CHANIR.WISMAIN\Downloads\gene_word_dataframe_activates.csv'

# Save the DataFrame to the specified file path
filtered_df.to_csv(file_path, index=False)
#gene_word_dataframe.to_csv('gene_word_dataframe.txt', index=False)
#suppresses


# In[511]:


#C:\Users\CHANIR.WISMAIN\Downloads\gene_word_dataframe_suppresses2.csv
#C:\Users\CHANIR.WISMAIN\Downloads\gene_word_dataframe_Stimulates2.csv
##new
#C:\Users\CHANIR.WISMAIN\Downloads\gene_word_dataframe_interacts4.cs
#C:\Users\CHANIR.WISMAIN\Downloads\gene_word_dataframe_Suppresses4.csv
#C:\Users\CHANIR.WISMAIN\Downloads\gene_word_dataframe_Stimulates4.cs
#C:\Users\CHANIR.WISMAIN\Downloads\gene_word_dataframe_binds4.csv'


# In[ ]:


file_path_activate1 = r'C:\Users\CHANIR.WISMAIN\Downloads\gene_word_dataframe_activates11.csv'
#'C:\Users\CHANIR.WISMAIN\Downloads\gene_word_dataframe_binds.csv'
file_path_interacts1 = r'C:\Users\CHANIR.WISMAIN\Downloads\gene_word_dataframe_interacts11.csv'
file_path_binds1 = r'C:\Users\CHANIR.WISMAIN\Downloads\gene_word_dataframe_binds11.csv'
parser_binds1 = pd.read_csv(file_path_binds1)
parser_activate1 = pd.read_csv(file_path_activate1)
parser_interacts1 = pd.read_csv(file_path_interacts1)


# In[ ]:


all_connected_tables = pd.concat([filtered_df,parser_binds1, parser_activate1],axis=0)
all_connected_tables


# In[ ]:


file_path_total = r'C:\Users\CHANIR.WISMAIN\Downloads\gene_word_dataframe_21_Jan_2024_3pm.csv'
all_connected_tables.to_csv(file_path_total , index=False)

